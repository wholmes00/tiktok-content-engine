"""
TikTok Content Engine — Edit Guide Generator (v3)
==================================================
Renders 10 video concepts (5 hero + 5 remix) as a clean .docx.

This is Document 1 in the pipeline. It provides everything an editor
needs to assemble each video: title, hook, duration, audio type,
full script, b-roll references, and on-screen text descriptions.

The Shoot Guide (Document 2) is derived FROM this data.

Usage:
    from templates.edit_guide_generator import generate_edit_guide
    generate_edit_guide(content, output_path)
"""

import os
from docx import Document

try:
    from templates.styles import (
        DARK, BLUE, GREEN, PURPLE, GRAY_TEXT, BODY_TEXT, RED,
        SIZE_EDIT_TITLE, SIZE_SUBTITLE, SIZE_SECTION_HEADER, SIZE_SUB_HEADER,
        SIZE_BODY_SMALL, SIZE_BODY_MEDIUM, SIZE_BODY, SIZE_EDIT_SECTION,
        SIZE_EDIT_HERO_TITLE, SIZE_TAG,
        FONT_PRIMARY, FONT_NARROW,
        add_centered_text, add_sub_header, add_blank, add_bullet_item,
        add_top_border, add_bottom_border, make_run, set_left_indent,
    )
except ImportError:
    from v2.templates.styles import (
        DARK, BLUE, GREEN, PURPLE, GRAY_TEXT, BODY_TEXT, RED,
        SIZE_EDIT_TITLE, SIZE_SUBTITLE, SIZE_SECTION_HEADER, SIZE_SUB_HEADER,
        SIZE_BODY_SMALL, SIZE_BODY_MEDIUM, SIZE_BODY, SIZE_EDIT_SECTION,
        SIZE_EDIT_HERO_TITLE, SIZE_TAG,
        FONT_PRIMARY, FONT_NARROW,
        add_centered_text, add_sub_header, add_blank, add_bullet_item,
        add_top_border, add_bottom_border, make_run, set_left_indent,
    )


def validate_content(content):
    """Validate content dict has all required fields."""
    required_top = ["creator_name", "product_name", "video_count",
                    "analysis_count", "date", "videos"]
    for key in required_top:
        if key not in content:
            raise ValueError(f"Missing required top-level key: '{key}'")

    for i, video in enumerate(content["videos"]):
        for field in ["number", "type", "title", "hook", "duration",
                      "audio", "script", "on_screen_text"]:
            if field not in video:
                raise ValueError(f"Video {i+1} missing field: '{field}'")
        for j, line in enumerate(video["script"]):
            if "type" not in line or "text" not in line:
                raise ValueError(f"Video {i+1}, script line {j+1} missing 'type' or 'text'")


def generate_edit_guide(content, output_path):
    """Generate an edit guide .docx from structured content.

    Args:
        content: dict matching the v3 schema (see CONTENT_SCHEMA)
        output_path: where to save the .docx file

    Returns:
        output_path on success
    """
    validate_content(content)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_PRIMARY

    # ── TITLE BLOCK ──
    add_centered_text(doc, content["creator_name"], SIZE_EDIT_TITLE, DARK, bold=True)
    add_centered_text(doc, content["product_name"], SIZE_SUBTITLE, GRAY_TEXT)

    p = add_centered_text(doc, "TikTok Affiliate Edit Guide",
                           SIZE_SECTION_HEADER, BLUE, bold=True)
    add_top_border(p)

    add_centered_text(doc, content["video_count"], SIZE_BODY_MEDIUM, GRAY_TEXT)
    add_centered_text(doc,
        f"Based on analysis of {content['analysis_count']} top-performing TikTok affiliate videos",
        SIZE_BODY_SMALL, GRAY_TEXT)
    add_centered_text(doc, content["date"], SIZE_BODY_SMALL, GRAY_TEXT)

    add_blank(doc)

    # ── HERO VIDEOS ──
    heroes = [v for v in content["videos"] if v["type"] == "hero"]
    remixes = [v for v in content["videos"] if v["type"] == "remix"]

    if heroes:
        add_centered_text(doc, "SECTION A", SIZE_BODY_SMALL, BLUE, bold=True)
        add_centered_text(doc, "Hero Videos", SIZE_EDIT_SECTION, DARK, bold=True)
        add_centered_text(doc, "On-camera storytelling — creator speaks directly to the audience",
                           SIZE_BODY_MEDIUM, GRAY_TEXT)

        for video in heroes:
            _render_video(doc, video)

    # ── REMIX VIDEOS ──
    if remixes:
        add_centered_text(doc, "SECTION B", SIZE_BODY_SMALL, BLUE, bold=True)
        add_centered_text(doc, "Remix Videos", SIZE_EDIT_SECTION, DARK, bold=True)
        add_centered_text(doc, "Short, music-driven edits — b-roll with voiceover or text only",
                           SIZE_BODY_MEDIUM, GRAY_TEXT)

        for video in remixes:
            _render_video(doc, video)

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def _render_video(doc, video):
    """Render a single video concept."""
    vtype = video["type"].upper()
    num = video["number"]

    # ── Header: type badge + title ──
    p = doc.add_paragraph()
    make_run(p, f"{vtype} VIDEO {num}", size=SIZE_BODY_SMALL, color=BLUE, bold=True)

    p = doc.add_paragraph()
    make_run(p, video["title"], size=SIZE_EDIT_HERO_TITLE, color=DARK, bold=True)
    add_bottom_border(p)

    # ── Meta line: hook ──
    p = doc.add_paragraph()
    make_run(p, "Hook: ", size=SIZE_BODY_MEDIUM, color=DARK, bold=True)
    make_run(p, f'"{video["hook"]}"', size=SIZE_BODY_MEDIUM, color=BODY_TEXT)

    # ── Meta line: duration + audio ──
    p = doc.add_paragraph()
    make_run(p, "Duration: ", size=SIZE_BODY_MEDIUM, color=DARK, bold=True)
    make_run(p, video["duration"], size=SIZE_BODY_MEDIUM, color=GRAY_TEXT)
    make_run(p, "    Audio: ", size=SIZE_BODY_MEDIUM, color=DARK, bold=True)
    make_run(p, video["audio"], size=SIZE_BODY_MEDIUM, color=GRAY_TEXT)

    add_blank(doc)

    # ── Script ──
    add_sub_header(doc, "SCRIPT")

    for i, line in enumerate(video["script"]):
        p = doc.add_paragraph()
        set_left_indent(p, 200)

        line_type = line["type"].upper()
        if "CAMERA" in line_type or line_type == "ON_CAMERA":
            tag_text = "ON CAMERA"
            tag_color = BLUE
        elif "VOICE" in line_type or line_type == "VOICEOVER":
            tag_text = "VOICEOVER"
            tag_color = PURPLE
        elif "DIRECTION" in line_type:
            tag_text = "DIRECTION"
            tag_color = GRAY_TEXT
        else:
            tag_text = line_type
            tag_color = GRAY_TEXT

        # Tag
        make_run(p, f"[{tag_text}]", size=SIZE_TAG, color=tag_color, bold=True)
        make_run(p, "  ", size=SIZE_TAG, color=tag_color)

        # Script text
        text_color = GRAY_TEXT if "DIRECTION" in line_type else BODY_TEXT
        make_run(p, line["text"], size=SIZE_BODY, color=text_color)

    add_blank(doc)

    # ── B-Roll Used (if any) ──
    broll_used = video.get("broll_used", [])
    if broll_used:
        add_sub_header(doc, "B-ROLL USED")
        for shot in broll_used:
            p = doc.add_paragraph()
            set_left_indent(p, 200)
            if isinstance(shot, dict):
                ts = shot.get("timestamp", "")
                code = shot.get("code", "")
                desc = shot.get("description", "")
                if ts:
                    make_run(p, f"{ts} ", font=FONT_NARROW, size=SIZE_BODY_SMALL,
                             color=GREEN, bold=True)
                make_run(p, code, size=SIZE_BODY, color=GREEN, bold=True)
                make_run(p, f"  {desc}", size=SIZE_BODY, color=BODY_TEXT)
            else:
                make_run(p, str(shot), size=SIZE_BODY, color=GREEN)
        add_blank(doc)

    # ── On-Screen Text ──
    add_sub_header(doc, "ON-SCREEN TEXT")
    for item in video["on_screen_text"]:
        p = doc.add_paragraph()
        set_left_indent(p, 200)
        if isinstance(item, dict):
            ts = item.get("timestamp", "")
            text = item.get("text", "")
            if ts:
                make_run(p, f"{ts} ", font=FONT_NARROW, size=SIZE_BODY_SMALL,
                         color=RED, bold=True)
            make_run(p, text, size=SIZE_BODY, color=BODY_TEXT)
        else:
            make_run(p, str(item), size=SIZE_BODY, color=BODY_TEXT)

    add_blank(doc)


# ═══════════════════════════════════════════════════════════════
# CONTENT SCHEMA (v3)
# ═══════════════════════════════════════════════════════════════

CONTENT_SCHEMA = {
    "creator_name": "str — e.g. MICHELLE'S FIRST SHOOT",
    "product_name": "str — full product name",
    "video_count": "str — e.g. 5 Hero Videos + 5 Remix Videos",
    "analysis_count": "int — number of videos in database",
    "date": "str — e.g. April 2026",

    "videos": [
        {
            "number": "int — 1-based index within type",
            "type": "str — 'hero' or 'remix'",
            "title": "str — creative title (e.g. The Swap Reveal)",
            "hook": "str — the hook line text",
            "duration": "str — e.g. ~30 seconds",
            "audio": "str — e.g. Original creator audio",
            "script": [
                {
                    "type": "str — 'on_camera', 'voiceover', or 'direction'",
                    "text": "str — the script line"
                }
            ],
            "broll_used": [
                {
                    "code": "str — B1, B3, etc.",
                    "description": "str — what the shot shows"
                }
            ],
            "on_screen_text": [
                "str — description of what text overlay to show"
            ]
        }
    ]
}
