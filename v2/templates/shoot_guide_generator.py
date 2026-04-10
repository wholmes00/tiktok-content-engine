"""
TikTok Content Engine — Shoot Guide Generator (v3)
===================================================
Renders a shooting checklist .docx DERIVED from the edit guide data.

This is Document 2 in the pipeline. It takes the 10 video concepts
from the edit guide and reorganizes everything into 3 clean sections
for the model/actress:

  1. ON-CAMERA LINES — grouped by video, just the lines to say
  2. B-ROLL SHOTS — deduplicated list of everything to film
  3. VOICEOVER RECORDINGS — grouped by video

The model doesn't need to know how videos are assembled — she just
needs to efficiently shoot all the content.

Usage:
    from templates.shoot_guide_generator import generate_shoot_guide
    generate_shoot_guide(content, output_path)

    # Or derive from edit guide data:
    from templates.shoot_guide_generator import derive_shoot_guide, generate_shoot_guide
    shoot_data = derive_shoot_guide(edit_guide_content)
    generate_shoot_guide(shoot_data, output_path)
"""

import os
from docx import Document

try:
    from templates.styles import (
        DARK, BLUE, GREEN, PURPLE, GRAY_TEXT, BODY_TEXT,
        SIZE_TITLE, SIZE_SUBTITLE, SIZE_BODY_SMALL, SIZE_BODY_MEDIUM,
        SIZE_SECTION_HEADER, SIZE_SUB_HEADER, SIZE_BODY, SIZE_TAG,
        SIZE_BROLL_LABEL,
        FONT_PRIMARY,
        add_centered_text, add_section_divider, add_blank,
        add_bottom_border, make_run, set_left_indent,
    )
except ImportError:
    from v2.templates.styles import (
        DARK, BLUE, GREEN, PURPLE, GRAY_TEXT, BODY_TEXT,
        SIZE_TITLE, SIZE_SUBTITLE, SIZE_BODY_SMALL, SIZE_BODY_MEDIUM,
        SIZE_SECTION_HEADER, SIZE_SUB_HEADER, SIZE_BODY, SIZE_TAG,
        SIZE_BROLL_LABEL,
        FONT_PRIMARY,
        add_centered_text, add_section_divider, add_blank,
        add_bottom_border, make_run, set_left_indent,
    )


def derive_shoot_guide(edit_guide_content):
    """Derive shoot guide data from edit guide content.

    Takes the v3 edit guide schema and extracts:
      - on_camera: all on-camera lines grouped by video
      - broll: deduplicated b-roll shot list
      - voiceovers: all voiceover scripts grouped by video

    Args:
        edit_guide_content: dict matching the edit guide v3 schema

    Returns:
        dict matching the shoot guide v3 schema
    """
    videos = edit_guide_content["videos"]

    on_camera_groups = []
    voiceover_groups = []
    broll_seen = {}  # code → description (dedup)

    for video in videos:
        # Use the title directly — it already includes "Hero N — " or "Remix N — "
        label = video["title"]

        if video["type"] == "hero":
            # HERO videos: ALL lines go in the on-camera section, in order,
            # tagged as ON CAMERA or VOICEOVER. The model records everything
            # in sequence — editors lay b-roll over the VO parts.
            tagged_lines = []
            for line in video["script"]:
                lt = line["type"].lower()
                if "voice" in lt or lt == "voiceover":
                    tag = "VOICEOVER"
                else:
                    tag = "ON CAMERA"
                tagged_lines.append({"tag": tag, "text": line["text"]})

            if tagged_lines:
                on_camera_groups.append({
                    "video_label": label,
                    "hook_template": video.get("hook_template", ""),
                    "lines": tagged_lines,
                })

        elif video["type"] == "remix":
            # REMIX videos: on-camera lines go in on-camera section (if any)
            on_camera_lines = []
            for line in video["script"]:
                lt = line["type"].lower()
                if "camera" in lt or lt == "on_camera":
                    on_camera_lines.append({"tag": "ON CAMERA", "text": line["text"]})

            if on_camera_lines:
                on_camera_groups.append({
                    "video_label": label,
                    "hook_template": video.get("hook_template", ""),
                    "lines": on_camera_lines,
                })

            # REMIX voiceovers go in the voiceover section
            vo_lines = []
            for line in video["script"]:
                lt = line["type"].lower()
                if "voice" in lt or lt == "voiceover":
                    vo_lines.append(line["text"])

            if vo_lines:
                voiceover_groups.append({
                    "video_label": label,
                    "script": " ".join(vo_lines),
                })

        # Collect b-roll (dedup by code)
        for shot in video.get("broll_used", []):
            if isinstance(shot, dict):
                code = shot.get("code", "")
                desc = shot.get("description", "")
                if code and code not in broll_seen:
                    broll_seen[code] = desc
            elif isinstance(shot, str):
                # Parse "B3 — description" format
                if " — " in shot:
                    parts = shot.split(" — ", 1)
                    code = parts[0].strip()
                    desc = parts[1].strip()
                elif " - " in shot:
                    parts = shot.split(" - ", 1)
                    code = parts[0].strip()
                    desc = parts[1].strip()
                else:
                    code = shot.strip()
                    desc = ""
                if code and code not in broll_seen:
                    broll_seen[code] = desc

    # Sort b-roll by code number
    def broll_sort_key(item):
        code = item["code"]
        try:
            return int(code.replace("B", "").replace("b", ""))
        except ValueError:
            return 999

    broll_list = [{"code": k, "description": v} for k, v in broll_seen.items()]
    broll_list.sort(key=broll_sort_key)

    total_videos = len(on_camera_groups) + len(voiceover_groups)

    return {
        "title": "SHOOT GUIDE",
        "subtitle": edit_guide_content.get("product_name", ""),
        "product_name": edit_guide_content.get("product_name", ""),
        "tagline": f"{len(edit_guide_content['videos'])} videos from one shoot — here's everything you need to capture",
        "on_camera": on_camera_groups,
        "broll": broll_list,
        "voiceovers": voiceover_groups,
    }


def validate_content(content):
    """Validate shoot guide content dict."""
    required = ["title", "subtitle", "on_camera", "broll", "voiceovers"]
    for key in required:
        if key not in content:
            raise ValueError(f"Missing required key: '{key}'")


def generate_shoot_guide(content, output_path):
    """Generate a shoot guide .docx from structured content.

    Args:
        content: dict matching CONTENT_SCHEMA (or output of derive_shoot_guide)
        output_path: where to save the .docx file

    Returns:
        output_path on success
    """
    validate_content(content)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_PRIMARY

    # ── TITLE BLOCK ──
    add_centered_text(doc, "SHOOT GUIDE", SIZE_TITLE, DARK, bold=True)
    add_centered_text(doc, content["subtitle"], SIZE_SUBTITLE, BLUE, bold=True)

    p = add_centered_text(doc, content.get("tagline",
        "Here's everything you need to capture"), SIZE_BODY_SMALL, GRAY_TEXT)
    add_bottom_border(p)

    # ── SECTION 1: ON-CAMERA LINES ──
    on_camera = content["on_camera"]
    total_lines = sum(len(g["lines"]) for g in on_camera)
    add_section_divider(doc, f"ON-CAMERA LINES ({len(on_camera)} videos, {total_lines} lines)", color=BLUE)

    p = doc.add_paragraph()
    make_run(p, "Film each line as its own take — ",
             size=SIZE_BODY_SMALL, color=GRAY_TEXT)
    make_run(p, "don't try to do it all in one shot.",
             size=SIZE_BODY_SMALL, color=GRAY_TEXT, bold=True)
    make_run(p, " The editor will stitch them together.",
             size=SIZE_BODY_SMALL, color=GRAY_TEXT)

    for group in on_camera:
        # Video header
        p = doc.add_paragraph()
        make_run(p, group["video_label"], size=SIZE_SUB_HEADER, color=DARK, bold=True)
        add_bottom_border(p, color=BLUE, size="4")

        # Script lines — each tagged with ON CAMERA or VOICEOVER
        for line in group["lines"]:
            p = doc.add_paragraph()
            set_left_indent(p, 200)

            if isinstance(line, dict):
                tag = line.get("tag", "ON CAMERA")
                text = line.get("text", "")
                tag_color = PURPLE if "VOICEOVER" in tag.upper() else BLUE
                make_run(p, f"{tag}   ", size=SIZE_TAG, color=tag_color, bold=True)
                make_run(p, text, size=SIZE_BODY, color=BODY_TEXT)
            else:
                # Legacy: plain string
                make_run(p, str(line), size=SIZE_BODY, color=BODY_TEXT)

        add_blank(doc)

    # ── SECTION 2: B-ROLL SHOTS ──
    broll = content["broll"]
    add_section_divider(doc, f"B-ROLL SHOTS ({len(broll)} clips)", color=GREEN)

    p = doc.add_paragraph()
    make_run(p, "Product shots, action shots, lifestyle moments. Shoot everything in natural light. "
             "Keep it simple — clean surface, steady hand.",
             size=SIZE_BODY_SMALL, color=GRAY_TEXT)

    add_blank(doc)

    for shot in broll:
        p = doc.add_paragraph()
        make_run(p, shot["code"], size=SIZE_BROLL_LABEL, color=GREEN, bold=True)
        make_run(p, f"  {shot['description']}", size=SIZE_BODY, color=BODY_TEXT)

    add_blank(doc)

    # ── SECTION 3: REMIX VOICEOVERS ──
    voiceovers = content["voiceovers"]
    add_section_divider(doc, f"REMIX VOICEOVERS ({len(voiceovers)} clips)", color=PURPLE)

    p = doc.add_paragraph()
    make_run(p, "These are audio-only voiceovers for the remix videos. Record in a quiet space — "
             "natural voice, don't perform it.",
             size=SIZE_BODY_SMALL, color=GRAY_TEXT)

    add_blank(doc)

    for vo in voiceovers:
        # Video header
        p = doc.add_paragraph()
        make_run(p, vo["video_label"], size=SIZE_SUB_HEADER, color=PURPLE, bold=True)

        # VO script
        p = doc.add_paragraph()
        set_left_indent(p, 200)
        run = p.add_run(f'"{vo["script"]}"')
        run.font.name = FONT_PRIMARY
        run.font.color.rgb = __import__('docx.shared', fromlist=['RGBColor']).RGBColor.from_string(PURPLE)
        run.bold = False

        add_blank(doc)

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


# ═══════════════════════════════════════════════════════════════
# CONTENT SCHEMA (v3)
# ═══════════════════════════════════════════════════════════════

CONTENT_SCHEMA = {
    "title": "SHOOT GUIDE",
    "subtitle": "str — product name",
    "product_name": "str — full product name",
    "tagline": "str — '10 videos from one shoot...'",

    "on_camera": [
        {
            "video_label": "str — e.g. Hero 1 — The Swap Reveal",
            "lines": ["str — each line to say on camera"]
        }
    ],

    "broll": [
        {
            "code": "str — B1, B2, etc.",
            "description": "str — what to shoot"
        }
    ],

    "voiceovers": [
        {
            "video_label": "str — e.g. Remix 1 — The Comparison",
            "script": "str — full VO script text"
        }
    ]
}
