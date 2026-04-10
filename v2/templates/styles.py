"""
TikTok Content Engine — Document Styling Constants
===================================================
Single source of truth for ALL document formatting.

Every color, font, size, and border spec is extracted directly from
the approved template documents (locked 2026-03-31). Both the shoot
guide and edit guide generators import from this file.

DO NOT hardcode styling values anywhere else. If it's a color, font,
or size — it lives here.

Source documents:
  - templates/approved/shoot_guide_template.docx
  - templates/approved/edit_guide_template.docx
"""

from docx.shared import Pt, Emu, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# ═══════════════════════════════════════════════════════════════
# COLORS (hex strings — used with RGBColor.from_string())
# ═══════════════════════════════════════════════════════════════

DARK = "1A1A2E"          # Primary dark text (titles, bold labels)
BLUE = "2563EB"          # Section headers, "ON CAMERA" tags, borders
GREEN = "059669"          # B-roll entries, timeline timestamps (edit guide)
PURPLE = "7C3AED"         # Voiceover tags and scripts (shoot guide)
GRAY_TEXT = "6B7280"      # Helper text, subtitles, bullet notes
BODY_TEXT = "333333"      # Main body text, script lines
RED = "DC2626"            # On-screen text timestamps (edit guide)


# ═══════════════════════════════════════════════════════════════
# FONTS
# ═══════════════════════════════════════════════════════════════

FONT_PRIMARY = "Arial"
FONT_NARROW = "Arial Narrow"   # Used for TH codes and timestamps


# ═══════════════════════════════════════════════════════════════
# FONT SIZES (in EMU — extracted from approved templates)
# ═══════════════════════════════════════════════════════════════

# --- Shared across both guides ---
SIZE_TITLE = Emu(279400)         # "SHOOT GUIDE" / top line (22pt)
SIZE_SUBTITLE = Emu(177800)      # Blue subtitle line (14pt)
SIZE_SECTION_HEADER = Emu(152400)  # "ON-CAMERA LINES", "B-ROLL SHOTS" (12pt)
SIZE_SUB_HEADER = Emu(139700)    # TH sub-headers, "TIMELINE", "ON-SCREEN TEXT" (11pt)
SIZE_BODY_SMALL = Emu(114300)    # Helper text, timestamps, small body (9pt)
SIZE_BODY = Emu(120650)          # Script text, list items (9.5pt)
SIZE_BODY_MEDIUM = Emu(127000)   # Hook/Audio labels (10pt)
SIZE_BROLL_LABEL = Emu(133350)   # B-roll "B1 PRODUCT" labels (10.5pt)
SIZE_BULLET_NOTE = Emu(107950)   # Bullet point notes under TH headers (8.5pt)
SIZE_TAG = Emu(101600)           # "ON CAMERA" / "VOICEOVER" tags (8pt)

# --- Edit guide specific ---
SIZE_EDIT_TITLE = Emu(330200)    # "MICHELLE'S FIRST SHOOT" (26pt)
SIZE_EDIT_SECTION = Emu(215900)  # "Hero Videos", "B-Roll Remix Videos" (17pt)
SIZE_EDIT_HERO_TITLE = Emu(190500)  # "Hero 1 — Name" (15pt)


# ═══════════════════════════════════════════════════════════════
# BORDER HELPERS
# ═══════════════════════════════════════════════════════════════

def add_bottom_border(paragraph, color=BLUE, size="8", space="4"):
    """Add a colored bottom border to a paragraph.

    Default: blue bottom border (used for TIMELINE, ON-SCREEN TEXT,
    B-ROLL ASSEMBLY, VOICEOVER SCRIPT sub-headers).
    """
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:color="{color}" w:sz="{size}" w:space="{space}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_top_border(paragraph, color=BLUE, size="4", space="8"):
    """Add a colored top border to a paragraph.

    Used on the "TikTok Affiliate Edit Guide" title line.
    """
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:color="{color}" w:sz="{size}" w:space="{space}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════
# INDENT HELPERS
# ═══════════════════════════════════════════════════════════════

def set_left_indent(paragraph, twips=400):
    """Set left indent on a paragraph (in twips).

    Default 400 twips — used for indented content lines in
    edit guide timelines.
    """
    pPr = paragraph._element.get_or_add_pPr()
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="{twips}"/>')
    pPr.append(ind)


# ═══════════════════════════════════════════════════════════════
# PARAGRAPH BUILDERS
# ═══════════════════════════════════════════════════════════════
# These create fully-styled paragraphs matching the approved
# templates exactly. Generators should use ONLY these functions
# to add content — never create raw paragraphs with inline styling.

def make_run(paragraph, text, font=FONT_PRIMARY, size=SIZE_BODY,
             color=BODY_TEXT, bold=False):
    """Add a styled run to an existing paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = size
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    return run


def add_centered_text(doc, text, size, color, bold=False, font=FONT_PRIMARY):
    """Add a center-aligned paragraph with a single styled run."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p, text, font=font, size=size, color=color, bold=bold)
    return p


def add_left_text(doc, text, size, color, bold=False, font=FONT_PRIMARY):
    """Add a left-aligned paragraph with a single styled run."""
    p = doc.add_paragraph()
    make_run(p, text, font=font, size=size, color=color, bold=bold)
    return p


def add_section_divider(doc, text, color=BLUE, size=SIZE_SECTION_HEADER):
    """Add a major section header with bottom border.

    Used for: ON-CAMERA LINES, B-ROLL SHOTS, VOICEOVER AUDIO (shoot guide)
    """
    p = doc.add_paragraph()
    make_run(p, text, size=size, color=color, bold=True)
    add_bottom_border(p, color=color)
    return p


def add_sub_header(doc, text, color=BLUE, size=SIZE_SUB_HEADER, border=True):
    """Add a sub-section header, optionally with bottom border.

    Used for: TIMELINE, ON-SCREEN TEXT, B-ROLL ASSEMBLY, VOICEOVER SCRIPT
    """
    p = doc.add_paragraph()
    make_run(p, text, size=size, color=color, bold=True)
    if border:
        add_bottom_border(p, color=color)
    return p


def add_bullet_item(doc, text, size=SIZE_BODY, color=BODY_TEXT):
    """Add a List Paragraph bullet item with visible bullet point."""
    p = doc.add_paragraph(style="List Paragraph")
    make_run(p, text, size=size, color=color)
    _ensure_bullet(p)
    return p


def add_blank(doc):
    """Add an empty paragraph (spacer)."""
    return doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# SHOOT GUIDE — SPECIFIC BUILDERS
# ═══════════════════════════════════════════════════════════════

def add_shoot_th_header(doc, code, product, title):
    """Add a talking-head sub-header with bottom border.

    Approved template format (runs):
      - Run 0: "TH" (blue, bold, 11pt)
      - Run 1: "1" (blue, bold, 11pt)
      - Run 2-3: "  Product — Title" (dark #1A1A2E, bold, inherit)

    The TH code is blue, the product + title are dark. Blue bottom border.
    """
    p = doc.add_paragraph()
    # Split the code (e.g. "TH1") to match approved run pattern
    make_run(p, code, font=FONT_PRIMARY, size=SIZE_SUB_HEADER,
             color=BLUE, bold=True)
    make_run(p, f"  {product} — {title}", font=FONT_PRIMARY,
             size=SIZE_SUB_HEADER, color=DARK, bold=True)
    add_bottom_border(p)
    return p


def add_shoot_script_line(doc, code, tag, script_text, is_direction=False):
    """Add a script line in the shoot guide format.

    Format: "ON CAMERA   'script text here'"
    Two runs:
      - Tag (Arial, blue or purple, bold, 8pt)
      - Script (Arial, body text or gray for directions, 9.5pt)

    Note: TH line codes (TH1-a, TH1-b) are intentionally omitted —
    only the tag (ON CAMERA / VOICEOVER) is shown.
    """
    p = doc.add_paragraph()

    # Run 1: Tag (ON CAMERA = blue, VOICEOVER = purple)
    tag_color = PURPLE if "VOICEOVER" in tag.upper() else BLUE
    make_run(p, f"{tag}   ", font=FONT_PRIMARY, size=SIZE_TAG,
             color=tag_color, bold=True)

    # Run 2: Script text or direction
    text_color = GRAY_TEXT if is_direction else BODY_TEXT
    make_run(p, script_text, font=FONT_PRIMARY, size=SIZE_BODY,
             color=text_color)

    return p


def add_shoot_broll_entry(doc, code, product, description):
    """Add a b-roll entry in the shoot guide format.

    Approved template format (4 runs):
      - Run 0-1: B-code only (green, bold, 10.5pt) — e.g. "B1"
      - Run 2-3: "  PRODUCT — description" (body text #333333, NO explicit size — inherits)

    Only the B-number is green. Product name and description are body text.
    Size is intentionally NOT set on the body runs to match the approved
    template's "inherit" pattern exactly.
    """
    p = doc.add_paragraph()
    make_run(p, code, font=FONT_PRIMARY,
             size=SIZE_BROLL_LABEL, color=GREEN, bold=True)
    # Do NOT set explicit size — approved template inherits here
    run = p.add_run(f"  {product} — {description}")
    run.font.name = FONT_PRIMARY
    run.font.color.rgb = RGBColor.from_string(BODY_TEXT)
    return p


def add_shoot_vo_entry(doc, code, product, script_text):
    """Add a voiceover entry in the shoot guide format.

    Approved template format: ALL purple, ALL bold throughout.
    Code+product runs have explicit 10.5pt size.
    Script runs inherit size (no explicit set) to match approved template.
    """
    p = doc.add_paragraph()
    make_run(p, f"{code}  {product} — ", font=FONT_PRIMARY,
             size=SIZE_BROLL_LABEL, color=PURPLE, bold=True)
    # Script: purple, NOT bold, NO explicit size — David requested unbolded scripts
    run = p.add_run(f'"{script_text}"')
    run.font.name = FONT_PRIMARY
    run.font.color.rgb = RGBColor.from_string(PURPLE)
    run.bold = False
    return p


def add_angle_callout(doc, content_angle, angle_evidence, angle_rationale=None):
    """Add a content angle callout line below a hero header.

    Renders as a compact data line showing the angle, rank, score, and
    video count — making it visible that this hero's concept is data-driven.

    Format: "Content Angle: shock_curiosity  |  #1 ranked  |  score 248,616  |  7 videos"
    Optional second line with rationale in gray italic.
    """
    if not content_angle or not angle_evidence:
        return None

    p = doc.add_paragraph()
    set_left_indent(p, 200)

    # Label
    make_run(p, "Content Angle: ", size=SIZE_BULLET_NOTE, color=BLUE, bold=True)

    # Angle name
    angle_display = content_angle.replace("_", " ").title()
    make_run(p, angle_display, size=SIZE_BULLET_NOTE, color=DARK, bold=True)

    # Separator + rank
    rank = angle_evidence.get("rank", "?")
    make_run(p, f"  |  ", size=SIZE_BULLET_NOTE, color=GRAY_TEXT)
    make_run(p, f"#{rank} ranked", size=SIZE_BULLET_NOTE, color=GREEN, bold=True)

    # Separator + score
    score = angle_evidence.get("weighted_score", 0)
    make_run(p, f"  |  ", size=SIZE_BULLET_NOTE, color=GRAY_TEXT)
    make_run(p, f"score {score:,}", size=SIZE_BULLET_NOTE, color=DARK)

    # Separator + video count
    count = angle_evidence.get("video_count", 0)
    make_run(p, f"  |  ", size=SIZE_BULLET_NOTE, color=GRAY_TEXT)
    make_run(p, f"{count} videos analyzed", size=SIZE_BULLET_NOTE, color=GRAY_TEXT)

    # Optional rationale line
    if angle_rationale:
        p2 = doc.add_paragraph()
        set_left_indent(p2, 200)
        run = p2.add_run(angle_rationale)
        run.font.name = FONT_PRIMARY
        run.font.size = SIZE_BULLET_NOTE
        run.font.color.rgb = RGBColor.from_string(GRAY_TEXT)
        run.italic = True
        return p2

    return p


def add_shoot_bullet_note(doc, text):
    """Add a bullet-point note under a TH header (shoot guide).

    Small gray text with visible bullet point.
    Uses List Paragraph style with explicit bullet numbering.
    """
    p = doc.add_paragraph(style="List Paragraph")
    make_run(p, text, size=SIZE_BULLET_NOTE, color=GRAY_TEXT)
    _ensure_bullet(p)
    return p


def _ensure_bullet(paragraph):
    """Add bullet numbering to a paragraph if not already present.

    This ensures List Paragraph items render with actual visible
    bullet points, not just indentation.
    """
    pPr = paragraph._element.get_or_add_pPr()
    # Check if numPr already exists
    existing = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
    if existing is not None:
        return
    numPr = parse_xml(
        f'<w:numPr {nsdecls("w")}>'
        f'<w:ilvl w:val="0"/>'
        f'<w:numId w:val="1"/>'
        f'</w:numPr>'
    )
    pPr.insert(0, numPr)


# ═══════════════════════════════════════════════════════════════
# EDIT GUIDE — SPECIFIC BUILDERS
# ═══════════════════════════════════════════════════════════════

def add_edit_hero_label(doc, text):
    """Add 'HERO VIDEO 1' label (blue, small, bold)."""
    p = doc.add_paragraph()
    make_run(p, text, size=SIZE_BODY_SMALL, color=BLUE, bold=True)
    return p


def add_edit_hero_title(doc, text):
    """Add 'Hero 1 — Title' (dark, 15pt, bold)."""
    p = doc.add_paragraph()
    make_run(p, text, size=SIZE_EDIT_HERO_TITLE, color=DARK, bold=True)
    return p


def add_edit_hook_line(doc, hook_text):
    """Add 'Hook: "text"' with bold label + body text."""
    p = doc.add_paragraph()
    make_run(p, "Hook: ", size=SIZE_BODY_MEDIUM, color=DARK, bold=True)
    make_run(p, f'"{hook_text}"', size=SIZE_BODY_MEDIUM, color=BODY_TEXT)
    return p


def add_edit_audio_line(doc, audio_text):
    """Add 'Audio: text' with bold label + gray text."""
    p = doc.add_paragraph()
    make_run(p, "Audio: ", size=SIZE_BODY_MEDIUM, color=DARK, bold=True)
    make_run(p, audio_text, size=SIZE_BODY_MEDIUM, color=GRAY_TEXT)
    return p


def add_edit_timeline_entry(doc, timestamp, shot_ref):
    """Add a timeline entry with green timestamp + shot reference.

    Format: "[0:00-0:03] TH1 CLOSE-UP"
    """
    p = doc.add_paragraph()
    make_run(p, f"{timestamp} ", font=FONT_NARROW, size=SIZE_BODY_SMALL,
             color=GREEN, bold=True)
    make_run(p, shot_ref, size=SIZE_BODY, color=GREEN, bold=True)
    return p


def add_edit_timeline_content(doc, text):
    """Add indented content under a timeline entry."""
    p = doc.add_paragraph()
    make_run(p, text, size=SIZE_BODY_SMALL, color=BODY_TEXT)
    set_left_indent(p, 400)
    return p


def add_edit_timeline_script(doc, text):
    """Add spoken script line under a timeline entry, indented and styled.

    Renders as italic blue text showing what the creator says at this point.
    Format: 'ON CAMERA: "script text here"' or 'VOICEOVER: "script text"'
    """
    p = doc.add_paragraph()
    make_run(p, text, size=SIZE_BODY_SMALL, color=BLUE)
    run = p.runs[0]
    run.italic = True
    set_left_indent(p, 400)
    return p


def add_edit_ost_entry(doc, timestamp, text):
    """Add on-screen text entry with red timestamp.

    Format: "[0:00-0:03] Text overlay content"
    """
    p = doc.add_paragraph()
    make_run(p, f"{timestamp} ", font=FONT_NARROW, size=SIZE_BODY_SMALL,
             color=RED, bold=True)
    make_run(p, text, size=SIZE_BODY, color=BODY_TEXT)
    return p


def add_edit_remix_title(doc, text):
    """Add 'Remix 1 — Title' (dark, 15pt, bold)."""
    p = doc.add_paragraph()
    make_run(p, text, size=SIZE_EDIT_HERO_TITLE, color=DARK, bold=True)
    return p


def add_edit_remix_info(doc, text):
    """Add product/type info line (gray, 9.5pt)."""
    p = doc.add_paragraph()
    make_run(p, text, size=SIZE_BODY, color=GRAY_TEXT)
    return p


def add_edit_body_text(doc, text):
    """Add a plain body text paragraph (9.5pt, body color)."""
    p = doc.add_paragraph()
    make_run(p, text, size=SIZE_BODY, color=BODY_TEXT)
    return p
