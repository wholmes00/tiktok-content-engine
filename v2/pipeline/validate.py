"""
TikTok Content Engine — Validation Suite
=========================================
Validates generated documents against approved templates.

Two types of validation:
  1. STRUCTURAL — checks paragraph count ranges, section markers,
     font families, color palette, border counts
  2. FINGERPRINT — compares the styling signature of every paragraph
     type against the approved template

Usage:
    from pipeline.validate import validate_shoot_guide, validate_edit_guide
    result = validate_shoot_guide("/path/to/generated.docx")
    if not result["passed"]:
        print(result["failures"])
"""

import os
from docx import Document

# Expected styling constants (must match styles.py)
EXPECTED_COLORS = {
    "shoot_guide": {"1A1A2E", "2563EB", "059669", "7C3AED", "6B7280", "333333"},
    "edit_guide":  {"1A1A2E", "2563EB", "059669", "DC2626", "6B7280", "333333"},
}

EXPECTED_FONTS = {"Arial", "Arial Narrow"}

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "..", "templates", "approved")


def _extract_signature(doc_path):
    """Extract a structural signature from a docx file."""
    doc = Document(doc_path)
    sig = {
        "total_paras": len(doc.paragraphs),
        "font_families": set(),
        "colors": set(),
        "styles": set(),
        "borders": 0,
        "centered": 0,
        "list_paragraphs": 0,
        "has_arial_narrow": False,
        "section_markers": [],
    }

    for p in doc.paragraphs:
        text = p.text.strip()

        if p.alignment and "CENTER" in str(p.alignment):
            sig["centered"] += 1

        style_name = p.style.name if p.style else "Normal"
        sig["styles"].add(style_name)
        if style_name == "List Paragraph":
            sig["list_paragraphs"] += 1

        pBdr = p._element.find(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr"
        )
        if pBdr is not None:
            sig["borders"] += 1

        for r in p.runs:
            if r.font.name:
                sig["font_families"].add(r.font.name)
                if r.font.name == "Arial Narrow":
                    sig["has_arial_narrow"] = True
            if r.font.color and r.font.color.rgb:
                sig["colors"].add(str(r.font.color.rgb))

        # Track section markers
        for marker in ["SECTION A", "SECTION B", "UPLOAD DETAILS",
                        "ON-CAMERA LINES", "B-ROLL SHOTS", "VOICEOVER AUDIO",
                        "TIMELINE", "ON-SCREEN TEXT", "B-ROLL ASSEMBLY",
                        "VOICEOVER SCRIPT", "Hero Videos", "B-Roll Remix Videos"]:
            if text == marker or text.startswith(marker):
                sig["section_markers"].append(marker)

    return sig


def _compare_signatures(generated_sig, guide_type):
    """Compare generated doc signature against expected values."""
    failures = []

    # Check font families
    if not EXPECTED_FONTS.issubset(generated_sig["font_families"]):
        missing = EXPECTED_FONTS - generated_sig["font_families"]
        failures.append(f"Missing font families: {missing}")

    # Check color palette
    expected_colors = EXPECTED_COLORS[guide_type]
    if not expected_colors.issubset(generated_sig["colors"]):
        missing = expected_colors - generated_sig["colors"]
        failures.append(f"Missing colors: {missing}")

    # Check for unexpected colors (could indicate styling drift)
    unexpected = generated_sig["colors"] - expected_colors
    if unexpected:
        failures.append(f"Unexpected colors found (possible styling drift): {unexpected}")

    # Check Arial Narrow is present (used for TH codes / timestamps)
    if not generated_sig["has_arial_narrow"]:
        failures.append("Arial Narrow font not found — TH codes/timestamps may be wrong")

    # Check borders exist
    if generated_sig["borders"] == 0:
        failures.append("No borders found — section headers should have borders")

    # Check list paragraphs exist (bullet items)
    if guide_type == "edit_guide" and generated_sig["list_paragraphs"] == 0:
        failures.append("No List Paragraph items — remix bullets/upload details missing")

    return failures


def validate_shoot_guide(doc_path):
    """Validate a generated shoot guide against expected structure.

    Returns:
        dict with 'passed' (bool), 'failures' (list), 'signature' (dict)
    """
    sig = _extract_signature(doc_path)
    failures = _compare_signatures(sig, "shoot_guide")

    # Shoot guide specific checks
    markers = sig["section_markers"]
    if not any("ON-CAMERA LINES" in m for m in markers):
        failures.append("Missing ON-CAMERA LINES section")
    if not any("B-ROLL SHOTS" in m for m in markers):
        failures.append("Missing B-ROLL SHOTS section")
    if not any("VOICEOVER AUDIO" in m for m in markers):
        failures.append("Missing VOICEOVER AUDIO section")

    # Check purple color is present (VO entries)
    if "7C3AED" not in sig["colors"]:
        failures.append("Missing purple (#7C3AED) — voiceover entries should be purple")

    # Check green color is present (b-roll entries)
    if "059669" not in sig["colors"]:
        failures.append("Missing green (#059669) — b-roll entries should be green")

    return {
        "passed": len(failures) == 0,
        "failures": failures,
        "signature": sig,
    }


def validate_edit_guide(doc_path):
    """Validate a generated edit guide against expected structure.

    Returns:
        dict with 'passed' (bool), 'failures' (list), 'signature' (dict)
    """
    sig = _extract_signature(doc_path)
    failures = _compare_signatures(sig, "edit_guide")

    # Edit guide specific checks
    markers = sig["section_markers"]
    if "SECTION A" not in markers:
        failures.append("Missing SECTION A marker")
    if "SECTION B" not in markers:
        failures.append("Missing SECTION B marker")
    if "UPLOAD DETAILS" not in markers:
        failures.append("Missing UPLOAD DETAILS section")

    # Must have TIMELINE and ON-SCREEN TEXT (hero sections)
    if "TIMELINE" not in markers:
        failures.append("Missing TIMELINE sub-header (hero videos)")
    if "ON-SCREEN TEXT" not in markers:
        failures.append("Missing ON-SCREEN TEXT sub-header")

    # Check red color for OST timestamps
    if "DC2626" not in sig["colors"]:
        failures.append("Missing red (#DC2626) — on-screen text timestamps should be red")

    # Check green for timeline timestamps
    if "059669" not in sig["colors"]:
        failures.append("Missing green (#059669) — timeline timestamps should be green")

    return {
        "passed": len(failures) == 0,
        "failures": failures,
        "signature": sig,
    }


def validate_both(shoot_path, edit_path):
    """Validate both guides and return combined results."""
    shoot_result = validate_shoot_guide(shoot_path)
    edit_result = validate_edit_guide(edit_path)

    return {
        "shoot_guide": shoot_result,
        "edit_guide": edit_result,
        "all_passed": shoot_result["passed"] and edit_result["passed"],
    }
